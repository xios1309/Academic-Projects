"""
Application de Creation de Documents avec Phrases et Images
-----------------------------------------------------------
Application Python avec interface graphique (Tkinter) qui genere automatiquement
des documents Word (.docx) a partir d'une liste de mots-cles.
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import os
import re
import random
import urllib.request
import threading
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- French sentence templates ---
SENTENCE_TEMPLATES = [
    "Le {mot} est vraiment magnifique aujourd'hui.",
    "J'ai decouvert un {mot} extraordinaire dans le jardin.",
    "Ma grand-mere adore le {mot} depuis son enfance.",
    "Le {mot} brille sous la lumiere du soleil.",
    "Nous avons observe un {mot} pendant notre promenade.",
    "Le petit enfant a dessine un {mot} sur son cahier.",
    "Il y a un {mot} au milieu de la place du village.",
    "Le professeur a parle du {mot} pendant le cours.",
    "On peut voir un {mot} depuis la fenetre de la cuisine.",
    "Le {mot} fait partie de notre vie quotidienne.",
    "Chaque matin, je contemple le {mot} avec admiration.",
    "Les enfants jouent pres du {mot} dans le parc.",
    "Un {mot} apparait soudainement devant nos yeux.",
    "La beaute du {mot} nous laisse sans voix.",
    "Tout le monde parle du {mot} dans le quartier.",
    "Le {mot} est un symbole de paix et d'harmonie.",
    "J'ai photographie un {mot} lors de mes vacances.",
    "Le {mot} se trouve au coeur de la foret.",
    "Mon ami m'a offert un {mot} pour mon anniversaire.",
    "Le {mot} represente la joie de vivre.",
    "Au lever du jour, le {mot} prend une teinte doree.",
    "Le {mot} est au centre de toutes les attentions.",
    "Personne n'a jamais vu un {mot} aussi impressionnant.",
    "Le {mot} fait rever les petits et les grands.",
    "On raconte que le {mot} porte bonheur.",
    "Le {mot} illumine la piece de sa presence.",
    "Nous avons installe un {mot} dans le salon.",
    "Le {mot} est le sujet de notre projet scolaire.",
    "Les touristes viennent admirer le {mot} chaque ete.",
    "Le {mot} est mentionne dans de nombreux livres.",
]


def generate_sentence(word, previous_sentences=None):
    """Generate a random French sentence containing the given word."""
    if previous_sentences is None:
        previous_sentences = []

    available = [t for t in SENTENCE_TEMPLATES if t not in previous_sentences]
    if not available:
        available = SENTENCE_TEMPLATES

    template = random.choice(available)
    sentence = template.format(mot=word)
    return sentence, template


class DocumentCreatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Createur de Documents - Phrases et Images")
        self.root.geometry("750x650")
        self.root.resizable(True, True)

        # Data
        self.words = []  # List of dicts: {word, sentence, template, image_path}
        self.image_style = tk.StringVar(value="photorealistic")
        self.cache_dir = "images_app"
        os.makedirs(self.cache_dir, exist_ok=True)

        self._build_ui()

    def _build_ui(self):
        """Build the complete user interface."""
        # --- Top frame: input ---
        frame_input = ttk.LabelFrame(self.root, text="Ajouter un mot-cle", padding=10)
        frame_input.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(frame_input, text="Mot :").grid(row=0, column=0, sticky=tk.W)
        self.entry_word = ttk.Entry(frame_input, width=30)
        self.entry_word.grid(row=0, column=1, padx=5)
        self.entry_word.bind("<Return>", lambda e: self.add_word())

        self.btn_add = ttk.Button(frame_input, text="Ajouter", command=self.add_word)
        self.btn_add.grid(row=0, column=2, padx=5)

        ttk.Label(frame_input, text="Style image :").grid(row=0, column=3, padx=(15, 5))
        style_combo = ttk.Combobox(
            frame_input,
            textvariable=self.image_style,
            values=["photorealistic", "illustration", "anime style"],
            state="readonly",
            width=15,
        )
        style_combo.grid(row=0, column=4)

        # --- Middle frame: word list ---
        frame_list = ttk.LabelFrame(self.root, text="Mots ajoutes", padding=10)
        frame_list.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        self.listbox = tk.Listbox(frame_list, height=10, font=("Arial", 11))
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.listbox.bind("<<ListboxSelect>>", self.on_select)

        scrollbar = ttk.Scrollbar(frame_list, orient=tk.VERTICAL, command=self.listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.listbox.config(yscrollcommand=scrollbar.set)

        # --- Buttons frame ---
        frame_buttons = ttk.Frame(self.root, padding=5)
        frame_buttons.pack(fill=tk.X, padx=10)

        self.btn_up = ttk.Button(frame_buttons, text="Monter", command=self.move_up)
        self.btn_up.pack(side=tk.LEFT, padx=2)

        self.btn_down = ttk.Button(frame_buttons, text="Descendre", command=self.move_down)
        self.btn_down.pack(side=tk.LEFT, padx=2)

        self.btn_new_phrase = ttk.Button(
            frame_buttons, text="Nouvelle phrase", command=self.regenerate_sentence
        )
        self.btn_new_phrase.pack(side=tk.LEFT, padx=2)

        self.btn_delete = ttk.Button(frame_buttons, text="Supprimer", command=self.delete_word)
        self.btn_delete.pack(side=tk.LEFT, padx=2)

        self.btn_generate = ttk.Button(
            frame_buttons, text="Generer le document", command=self.generate_document
        )
        self.btn_generate.pack(side=tk.RIGHT, padx=2)

        # --- Preview frame ---
        frame_preview = ttk.LabelFrame(self.root, text="Apercu de la phrase", padding=10)
        frame_preview.pack(fill=tk.X, padx=10, pady=5)

        self.label_preview = ttk.Label(
            frame_preview, text="(selectionnez un mot pour voir la phrase)", wraplength=700
        )
        self.label_preview.pack(fill=tk.X)

        # --- Progress frame ---
        frame_progress = ttk.Frame(self.root, padding=5)
        frame_progress.pack(fill=tk.X, padx=10, pady=5)

        self.progress = ttk.Progressbar(frame_progress, mode="determinate")
        self.progress.pack(fill=tk.X, side=tk.LEFT, expand=True, padx=(0, 10))

        self.label_status = ttk.Label(frame_progress, text="Pret")
        self.label_status.pack(side=tk.RIGHT)

    def add_word(self):
        """Add a new keyword to the list."""
        word = self.entry_word.get().strip()
        if not word:
            messagebox.showwarning("Attention", "Veuillez saisir un mot.")
            return

        # Check duplicate
        if any(item["word"].lower() == word.lower() for item in self.words):
            messagebox.showwarning("Attention", f"Le mot '{word}' est deja dans la liste.")
            return

        sentence, template = generate_sentence(word)
        entry = {
            "word": word,
            "sentence": sentence,
            "template": template,
            "image_path": None,
            "used_templates": [template],
        }
        self.words.append(entry)
        self.listbox.insert(tk.END, f"{word} - {sentence}")
        self.entry_word.delete(0, tk.END)

        # Select the newly added item
        self.listbox.selection_clear(0, tk.END)
        self.listbox.selection_set(tk.END)
        self.on_select(None)

    def on_select(self, event):
        """Update the preview when a word is selected."""
        sel = self.listbox.curselection()
        if sel:
            idx = sel[0]
            entry = self.words[idx]
            self.label_preview.config(text=f"Phrase : {entry['sentence']}")

    def move_up(self):
        """Move selected word up in the list."""
        sel = self.listbox.curselection()
        if not sel or sel[0] == 0:
            return
        idx = sel[0]
        self.words[idx], self.words[idx - 1] = self.words[idx - 1], self.words[idx]
        self._refresh_listbox()
        self.listbox.selection_set(idx - 1)
        self.on_select(None)

    def move_down(self):
        """Move selected word down in the list."""
        sel = self.listbox.curselection()
        if not sel or sel[0] == len(self.words) - 1:
            return
        idx = sel[0]
        self.words[idx], self.words[idx + 1] = self.words[idx + 1], self.words[idx]
        self._refresh_listbox()
        self.listbox.selection_set(idx + 1)
        self.on_select(None)

    def delete_word(self):
        """Delete selected word from the list."""
        sel = self.listbox.curselection()
        if not sel:
            return
        idx = sel[0]
        self.words.pop(idx)
        self._refresh_listbox()
        self.label_preview.config(text="(selectionnez un mot pour voir la phrase)")

    def regenerate_sentence(self):
        """Generate a new sentence for the selected word."""
        sel = self.listbox.curselection()
        if not sel:
            messagebox.showwarning("Attention", "Selectionnez un mot dans la liste.")
            return
        idx = sel[0]
        entry = self.words[idx]
        sentence, template = generate_sentence(entry["word"], entry["used_templates"])
        entry["sentence"] = sentence
        entry["template"] = template
        entry["used_templates"].append(template)
        self._refresh_listbox()
        self.listbox.selection_set(idx)
        self.label_preview.config(text=f"Phrase : {sentence}")

    def _refresh_listbox(self):
        """Refresh the listbox display."""
        self.listbox.delete(0, tk.END)
        for item in self.words:
            self.listbox.insert(tk.END, f"{item['word']} - {item['sentence']}")

    def download_image(self, word, style):
        """Download an AI-generated image from Pollinations.ai."""
        prompt = f"{word} {style}"
        encoded_prompt = urllib.request.quote(prompt)
        url = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width=512&height=512&nologo=true"

        filename = f"{word.replace(' ', '_')}_{style.replace(' ', '_')}.png"
        filepath = os.path.join(self.cache_dir, filename)

        # Use cached image if available
        if os.path.exists(filepath):
            return filepath

        try:
            urllib.request.urlretrieve(url, filepath)
            return filepath
        except Exception as e:
            print(f"Erreur lors du telechargement de l'image pour '{word}': {e}")
            return None

    def generate_document(self):
        """Generate the Word document with all words."""
        if not self.words:
            messagebox.showwarning("Attention", "Ajoutez au moins un mot avant de generer.")
            return

        # Ask for save location
        output_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Document Word", "*.docx")],
            title="Enregistrer le document",
        )
        if not output_path:
            return

        # Disable buttons during generation
        self.btn_generate.config(state=tk.DISABLED)
        self.progress["value"] = 0
        self.progress["maximum"] = len(self.words)

        # Run in background thread
        thread = threading.Thread(target=self._generate_worker, args=(output_path,))
        thread.daemon = True
        thread.start()

    def _generate_worker(self, output_path):
        """Worker thread for document generation."""
        try:
            # Try to use template if available
            template_path = os.path.join(os.path.dirname(__file__), "v9.docx")
            if os.path.exists(template_path):
                doc = Document(template_path)
                # Clear existing content
                for paragraph in doc.paragraphs:
                    p_element = paragraph._element
                    p_element.getparent().remove(p_element)
            else:
                doc = Document()

            style = self.image_style.get()
            total = len(self.words)

            for i, entry in enumerate(self.words):
                word = entry["word"]
                sentence = entry["sentence"]

                self.root.after(0, self._update_status, f"Traitement de '{word}' ({i+1}/{total})...")

                # Download image
                image_path = self.download_image(word, style)
                entry["image_path"] = image_path

                # Add image to document
                if image_path and os.path.exists(image_path):
                    paragraph_img = doc.add_paragraph()
                    paragraph_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = paragraph_img.add_run()
                    run_img.add_picture(image_path, width=Inches(5.5))

                # Add sentence with keyword highlighted in red
                paragraph_text = doc.add_paragraph()
                paragraph_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Split sentence around the keyword (case-insensitive)
                pattern = re.compile(re.escape(word), re.IGNORECASE)
                parts = pattern.split(sentence)
                matches = pattern.findall(sentence)

                for j, part in enumerate(parts):
                    if part:
                        run = paragraph_text.add_run(part)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(36)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                    if j < len(matches):
                        run_keyword = paragraph_text.add_run(matches[j])
                        run_keyword.font.name = "Times New Roman"
                        run_keyword.font.size = Pt(36)
                        run_keyword.font.bold = True
                        run_keyword.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

                # Add page break (except for last word)
                if i < total - 1:
                    doc.add_page_break()

                # Update progress
                self.root.after(0, self._update_progress, i + 1)

            # Save document
            doc.save(output_path)
            self.root.after(
                0, self._generation_complete, f"Document genere avec succes !\n{output_path}"
            )

        except Exception as e:
            self.root.after(0, self._generation_error, str(e))

    def _update_status(self, text):
        """Update status label."""
        self.label_status.config(text=text)

    def _update_progress(self, value):
        """Update progress bar."""
        self.progress["value"] = value

    def _generation_complete(self, message):
        """Handle successful generation."""
        self.btn_generate.config(state=tk.NORMAL)
        self.label_status.config(text="Termine !")
        messagebox.showinfo("Succes", message)

    def _generation_error(self, error_msg):
        """Handle generation error."""
        self.btn_generate.config(state=tk.NORMAL)
        self.label_status.config(text="Erreur")
        messagebox.showerror("Erreur", f"Erreur lors de la generation :\n{error_msg}")


if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentCreatorApp(root)
    root.mainloop()
